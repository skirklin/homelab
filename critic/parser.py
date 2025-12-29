"""
Document parser - handles .docx, .txt, and .md files.
"""

import re
from pathlib import Path
from typing import Optional, Union

from critic.types import ParsedDocument, Heading


def parse_document(
    file_path: Optional[Union[str, Path]] = None,
    content: Optional[str] = None,
    buffer: Optional[bytes] = None,
) -> ParsedDocument:
    """
    Parse a document from file path, string content, or bytes buffer.

    Args:
        file_path: Path to .docx, .txt, or .md file
        content: Raw text content
        buffer: Bytes buffer (for .docx files)

    Returns:
        ParsedDocument with paragraphs, headings, and full text
    """
    if file_path:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == '.docx':
            return _parse_docx(path)
        elif suffix in ('.txt', '.md', '.markdown'):
            text = path.read_text(encoding='utf-8')
            return parse_text(text, title=path.stem)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    elif buffer:
        return _parse_docx_buffer(buffer)

    elif content:
        return parse_text(content)

    else:
        raise ValueError("Must provide file_path, content, or buffer")


def parse_text(text: str, title: str = "Untitled") -> ParsedDocument:
    """
    Parse plain text or markdown content.

    Args:
        text: The text content
        title: Document title (defaults to first heading or 'Untitled')

    Returns:
        ParsedDocument
    """
    # Split into paragraphs (double newline separated)
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]

    # Find headings
    headings: list[Heading] = []

    for i, para in enumerate(paragraphs):
        # Check each line within the paragraph for headings
        # (paragraphs may contain multiple consecutive heading lines)
        lines = para.split('\n')
        found_heading = False

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Markdown headings
            md_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if md_match:
                level = len(md_match.group(1))
                heading_text = md_match.group(2).strip()
                headings.append(Heading(text=heading_text, level=level, paragraph_index=i))
                found_heading = True
                continue

            # All-caps lines that look like chapter headings
            if line.isupper() and len(line) < 100:
                headings.append(Heading(text=line, level=1, paragraph_index=i))
                found_heading = True
                continue

            # "Chapter X" or "CHAPTER X" patterns
            chapter_match = re.match(
                r'^(Chapter|CHAPTER)\s+(\d+|[IVXLC]+)\.?\s*(.*?)$',
                line,
                re.IGNORECASE
            )
            if chapter_match:
                headings.append(Heading(text=line, level=1, paragraph_index=i))
                found_heading = True
                continue

        if found_heading:
            continue

    # Try to find a proper title
    doc_title = title

    # First check if there's a title-like line at the start of the document
    # (common in Gutenberg texts: "Title\n\nby Author")
    first_lines = text.strip().split('\n')
    if first_lines:
        potential_title = first_lines[0].strip()
        # A good title candidate: short, not a chapter marker, not all-caps single word
        if (potential_title and
            len(potential_title) < 80 and
            not re.match(r'^(Chapter|CHAPTER)\s+(\d+|[IVXLC]+)', potential_title, re.IGNORECASE) and
            not potential_title.startswith('#')):
            doc_title = potential_title

    # Fall back to first heading if it's not a chapter marker
    if doc_title == title and headings:
        first_heading = headings[0].text
        if not re.match(r'^(Chapter|CHAPTER)\s+(\d+|[IVXLC]+)', first_heading, re.IGNORECASE):
            doc_title = first_heading

    return ParsedDocument(
        title=doc_title,
        paragraphs=paragraphs,
        headings=headings,
        full_text=text,
    )


def _parse_docx(path: Path) -> ParsedDocument:
    """Parse a .docx file."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for .docx parsing: pip install python-docx")

    doc = Document(str(path))
    paragraphs: list[str] = []
    headings: list[Heading] = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        paragraphs.append(text)
        para_index = len(paragraphs) - 1

        # Check for heading style
        style = para.style
        if style and style.name:
            style_name = style.name.lower()
            if 'heading' in style_name:
                # Extract level from style name like "Heading 1", "Heading 2"
                level_match = re.search(r'(\d+)', style_name)
                level = int(level_match.group(1)) if level_match else 1
                headings.append(Heading(text=text, level=level, paragraph_index=para_index))
            elif style_name == 'title':
                headings.append(Heading(text=text, level=1, paragraph_index=para_index))

    full_text = '\n\n'.join(paragraphs)

    # Determine title
    title = path.stem
    if headings:
        title = headings[0].text

    return ParsedDocument(
        title=title,
        paragraphs=paragraphs,
        headings=headings,
        full_text=full_text,
    )


def _parse_docx_buffer(buffer: bytes) -> ParsedDocument:
    """Parse a .docx file from bytes buffer."""
    import io
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for .docx parsing: pip install python-docx")

    doc = Document(io.BytesIO(buffer))
    paragraphs: list[str] = []
    headings: list[Heading] = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        paragraphs.append(text)
        para_index = len(paragraphs) - 1

        style = para.style
        if style and style.name:
            style_name = style.name.lower()
            if 'heading' in style_name:
                level_match = re.search(r'(\d+)', style_name)
                level = int(level_match.group(1)) if level_match else 1
                headings.append(Heading(text=text, level=level, paragraph_index=para_index))
            elif style_name == 'title':
                headings.append(Heading(text=text, level=1, paragraph_index=para_index))

    full_text = '\n\n'.join(paragraphs)
    title = headings[0].text if headings else "Untitled"

    return ParsedDocument(
        title=title,
        paragraphs=paragraphs,
        headings=headings,
        full_text=full_text,
    )
